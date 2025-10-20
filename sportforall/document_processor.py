# sportforall/document_processor.py

from docx import Document
from docx.shared import Inches, Pt # Для стилів таблиці, якщо потрібно
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT # Для вирівнювання тексту
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT # Для вирівнювання таблиці/комірок
from docx.table import _Cell, Table # Додано _Cell та Table напряму
from docx.text.paragraph import Paragraph # Додано Paragraph напряму
from docx.text.run import Run # Додано Run напряму

# Для роботи з XML елементами напряму - необхідно для вставки/видалення елементів
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl

import re
import os
import math
from copy import deepcopy # Для копіювання XML елементів

# Імпортуємо моделі даних та логер помилок
from sportforall.models import Contract, Item
from sportforall import error_handling
from sportforall import constants

# Регулярний вираз для пошуку плейсхолдерів <...>
PLACEHOLDER_PATTERN = re.compile(r"<([^>]+)>")


def find_placeholders_in_template(filepath: str) -> list:
    """
    Читає файл шаблону .docm або .docx та знаходить усі унікальні плейсхолдери <...>.
    Ця версія надійно шукає плейсхолдери, що охоплюють кілька run-ів,
    збираючи текст з усіх run-ів елемента перед пошуком.

    Args:
        filepath: Шлях до файлу шаблону.

    Returns:
        Список унікальних знайдених плейсхолдерів (включаючи дужки).
        Повертає порожній список у разі помилки читання файлу.
    """
    if not os.path.exists(filepath):
        print(f"Помилка: Файл шаблону не знайдено: {filepath}")
        return []

    placeholders = set()
    try:
        document = Document(filepath)

        # Функція для пошуку в елементі (параграфі або комірці)
        def search_in_element(element):
            # Якщо це параграф
            if isinstance(element, Paragraph): # <-- ЭТА ПРОВЕРКА КОРРЕКТНА здесь
                # Збираємо текст з усіх run-ів та їх початкові позиції
                full_text = ""
                for run in element.runs:
                    full_text += run.text

                # Шукаємо плейсхолдери у повному тексті
                for match in PLACEHOLDER_PATTERN.finditer(full_text):
                    placeholder = match.group(0) # Знайдений плейсхолдер з дужками
                    placeholders.add(placeholder)

            # Якщо це комірка, шукаємо в її параграфах
            elif isinstance(element, _Cell): # <-- ЭТА ПРОВЕРКА КОРРЕКТНА здесь
                 for para in element.paragraphs:
                     search_in_element(para)


        # Проходимо по всіх елементах body, щоб знайти всі параграфи та таблиці
        # Використовуємо ітератор по елементах body, щоб знайти всі параграфи та таблиці
        # в основному тексті, колонтитулах, зносках тощо.
        # Для простоти, наразі обмежуємось body
        # for element in document.element.body.iter(): # <-- ЗАКОММЕНТИРОВАТЬ ЭТОТ БЛОК
        #     # if isinstance(element, qn('w:p')): # Параграф
        #     #     search_in_element(Paragraph(element, document._body))
        #     # elif isinstance(element, qn('w:tc')): # Комірка таблиці
        #     #     search_in_element(_Cell(element, None)) # None замість батьківського row

        # Більш простий і безпечний спосіб для пошуку в основному тексті та таблицях Body:
        # УБЕДИТЕСЬ, ЧТО ЭТИ ЦИКЛЫ РАСКОММЕНТИРОВАНЫ:
        for para in document.paragraphs:
             search_in_element(para)

        for table in document.tables:
             for row in table.rows:
                  for cell in row.cells:
                       search_in_element(cell)

        print(f"Знайдено {len(placeholders)} унікальних плейсхолдерів у {filepath}")
        # print(f"Знайдені плейсхолдери: {sorted(list(placeholders))}") # Для відладки
        return sorted(list(placeholders))

    except Exception as e:
        error_logger.log_error(e, f"Помилка при читанні файлу шаблону або пошуку плейсхолдерів: {filepath}")
        print(f"Помилка при читанні файлу шаблону: {e}")
        return []


def _replace_placeholder_in_element(element, placeholder: str, replacement: str):
    """
    Надійно замінює текст плейсхолдера у параграфі або комірці таблиці,
    зберігаючи форматування.

    Args:
        element: Об'єкт Paragraph або _Cell.
        placeholder: Плейсхолдер для заміни (напр., "<адреса>").
        replacement: Текст, на який потрібно замінити плейсхолдер.
    """
    # Якщо це комірка, викликаємо рекурсивно для її параграфів
    if isinstance(element, _Cell):
        for para in element.paragraphs:
            _replace_placeholder_in_element(para, placeholder, replacement)
        return

    # Якщо це параграф
    if isinstance(element, Paragraph):
        original_text = element.text
        if placeholder not in original_text:
            return # Плейсхолдера немає в цьому параграфі

        # Використовуємо re.search для пошуку першого входження
        match = re.search(re.escape(placeholder), original_text)
        if match:
            start_index = match.start()
            end_index = match.end()

            # Визначаємо run-и, які охоплює плейсхолдер
            runs = element.runs
            run_starts = []
            current_pos = 0
            for run in runs:
                 run_starts.append(current_pos)
                 current_pos += len(run.text)

            # Знаходимо індекс першого run, що перетинається з плейсхолдером
            start_run_index = -1
            for i, run_start in enumerate(run_starts):
                 if run_start + len(runs[i].text) > start_index:
                      start_run_index = i
                      break

            if start_run_index == -1:
                 # Це не повинно статись, якщо match знайдено в original_text
                 print(f"Попередження: Не знайдено початковий run для плейсхолдера {placeholder}.")
                 # Виконуємо просту текстову заміну як запасний варіант
                 element.text = original_text.replace(placeholder, replacement)
                 return


            # Знаходимо індекс останнього run, що перетинається з плейсхолдером
            end_run_index = -1
            for i in range(len(runs) - 1, -1, -1): # Йдемо зі зворотного порядку
                 run_start = run_starts[i]
                 if run_start < end_index:
                      end_run_index = i
                      break

            if end_run_index == -1:
                 # Це не повинно статись
                 print(f"Попередження: Не знайдено кінцевий run для плейсхолдера {placeholder}.")
                 # Виконуємо просту текстову заміну як запасний варіант
                 element.text = original_text.replace(placeholder, replacement)
                 return


            # Отримуємо форматування першого run, який містить частину плейсхолдера
            # (будемо копіювати це форматування)
            first_run_in_placeholder = runs[start_run_index]
            # Зберігаємо властивості форматування (більш повноцінно)
            rpr = first_run_in_placeholder._r.get_or_add_rPr() # Отримуємо або додаємо властивості run
            # deepcopy rpr, якщо потрібно бути впевненим, що зміни не вплинуть на оригінал
            # rpr_copy = deepcopy(rpr)


            # Текст до плейсхолдера в першому run
            text_before_in_first_run = first_run_in_placeholder.text[:start_index - run_starts[start_run_index]]
            # Текст після плейсхолдера в останньому run
            text_after_in_last_run = runs[end_run_index].text[end_index - run_starts[end_run_index]:]


            # Очищаємо всі run-и, які повністю або частково містять плейсхолдер
            # Ітерація у зворотньому порядку безпечніша при видаленні/очищенні run-ів
            for i in range(end_run_index, start_run_index - 1, -1):
                 # element._element.remove(runs[i]._element) # Видалення елемента run (складно)
                 # Простіше: очистити текст run-а
                 runs[i].text = ""


            # Вставляємо текст до плейсхолдера в першому run
            runs[start_run_index].text = text_before_in_first_run

            # Вставляємо замінюючий текст як новий run після start_run_index
            new_run_element = OxmlElement('w:r') # Створюємо новий XML елемент run
            # Копіюємо властивості форматування з першого run
            new_run_element.append(rpr) # Додаємо скопійовані властивості форматування
            t_element = OxmlElement('w:t') # Створюємо XML елемент тексту
            t_element.text = replacement
            # Важливо: Встановлюємо xml:space="preserve", якщо замінюючий текст починається/закінчується пробілами
            if replacement.startswith(' ') or replacement.endswith(' '):
                 t_element.set(qn('xml:space'), 'preserve')
            new_run_element.append(t_element)

            # Вставляємо новий XML елемент run після XML елемента початкового run
            runs[start_run_index]._element.addnext(new_run_element)

            # Вставляємо текст після плейсхолдера в останньому run
            # Знаходимо останній run (він може бути іншим, якщо плейсхолдер охопив кілька run-ів)
            # Після вставки нового run з заміною, runs list оновлюється (або потрібно його оновити)
            # Щоб знайти, куди вставити текст після плейсхолдера, потрібно знати, де знаходився end_run_index
            # в оригінальній послідовності і вставити текст після нового замінюючого run.

            # Простіший підхід: створити новий run для тексту після плейсхолдера
            if text_after_in_last_run:
                 # Знаходимо елемент останнього run (або елемент, що йшов після нього)
                 # Це може бути складно, якщо run-и були видалені/очищені
                 # Найпростіше - додати його в кінець параграфа з форматуванням останнього run
                 # Отримуємо форматування останнього run, що містив частину плейсхолдера
                 # last_run_in_placeholder = runs[end_run_index] # Це вже не коректний індекс після модифікацій
                 # Спробуємо скопіювати форматування з того ж rpr, що і для замінюючого тексту
                 run_after = element.add_run(text_after_in_last_run)
                 run_after._r.append(deepcopy(rpr)) # Копіюємо властивості


            # print(f"Заміна в параграфі ({placeholder} -> {replacement}) - покращена")


        else:
             # Якщо placeholder не знайдено в original_text (хоча виклик був)
             pass # Нічого не робити


    # TODO: Ця функція все ще може мати проблеми зі складним форматуванням та структурою run-ів.
    # Найбільш надійний спосіб - працювати повністю з XML (lxml), що виходить за рамки простого використання python-docx.

def _create_items_table(items: list, contract: Contract) -> Table:
    """
    Створює об'єкт docx.table.Table з даними про товари та заголовками.
    Створюється в тимчасовому документі.

    Args:
        items: Список об'єктів Item.
        contract: Об'єкт Contract (потрібен для отримання загальної суми з полів).

    Returns:
        Об'єкт docx.table.Table.
    """
    # Визначаємо кількість колонок
    num_cols = 4 # Назва, ДК, Кількість, Сума

    # Створюємо таблицю в тимчасовому документі. 0 рядків спочатку.
    doc_temp = Document()
    table = doc_temp.add_table(rows=0, cols=num_cols)

    # Налаштування стилю таблиці
    table.style = 'Table Grid' # Застосовуємо базовий стиль

    # Додаємо рядок заголовків
    header_cells = table.add_row().cells
    header_cells[0].text = constants.ITEM_TABLE_HEADERS[0] # "Назва Товару"
    header_cells[1].text = constants.ITEM_TABLE_HEADERS[1] # "ДК 021:2015"
    header_cells[2].text = constants.ITEM_TABLE_HEADERS[2] # "Кількість"
    header_cells[3].text = constants.ITEM_TABLE_HEADERS[3] # "Сума, грн"


    # Налаштування стилю заголовків (жирний, вирівнювання по центру, вертикально по центру)
    for cell in header_cells:
         cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
         for p in cell.paragraphs:
              p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
              for run in p.runs:
                   run.bold = True

    # Додаємо рядки з даними товарів
    for item in items:
         row_cells = table.add_row().cells
         row_cells[0].text = item.name
         row_cells[1].text = item.dk
         # Форматуємо числові значення для відображення з комою
         row_cells[2].text = str(item.quantity).replace('.', ',') # Кількість
         row_cells[3].text = f"{item.total_sum:.2f}".replace('.', ',') # Сума з 2 знаками після коми та комою як роздільником

         # Вирівнювання для числових колонок (Кількість, Сума) по правому краю
         for i in [2, 3]:
              for p in row_cells[i].paragraphs:
                   p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

         # Вирівнювання тексту в комірках по вертикалі (по центру)
         for cell in row_cells:
              cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


    # Додаємо рядок "Всього"
    if items: # Додаємо рядок "Всього", тільки якщо є товари
        total_row = table.add_row()
        total_cells = total_row.cells

        # Визначаємо комірку для тексту "Всього" (перша комірка)
        total_cells[0].text = constants.ITEM_TOTAL_ROW_TEXT # "Всього:"

        # Об'єднуємо комірки для тексту "Всього"
        try:
             # Об'єднуємо першу комірку з наступними до передостанньої включно
             # Напр., для 4 колонок, об'єднуємо cell(0) з cell(1), cell(2)
             # Об'єднуємо total_cells[0] до total_cells[num_cols - 2]
             if num_cols > 1:
                  for i in range(1, num_cols - 1):
                       total_cells[0].merge(total_cells[i])


        except Exception as e:
             error_logger.log_error(e, "Помилка при об'єднанні комірок в рядку 'Всього' таблиці товарів")
             print(f"Помилка об'єднання комірок в таблиці товарів: {e}")


        # Вирівнювання тексту "Всього" по правому краю в об'єднаній комірці
        for p in total_cells[0].paragraphs:
             p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
             # Зробити текст "Всього" жирним
             for run in p.runs:
                  run.bold = True

        # Вставляємо загальну суму
        # Беремо значення з contract.fields, яке розрахував GUI (воно вже відформатоване з комою)
        total_sum_text = contract.fields.get(constants.TOTAL_SUM_PLACEHOLDER, "")
        # Знаходимо комірку, куди потрібно вставити загальну суму (остання комірка)
        total_cells[-1].text = total_sum_text

        # Вирівнювання загальної суми по правому краю
        for p in total_cells[-1].paragraphs:
             p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
             # Зробити текст суми жирним
             for run in p.runs:
                  run.bold = True

        # Вирівнювання тексту в комірках по вертикалі (по центру)
        for cell in total_cells:
              cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


    return table # Повертаємо створений об'єкт таблиці (з тимчасового документа)


def generate_document(contract: Contract, template_path: str, output_dir: str) -> str:
    """
    Генерує документ на основі шаблону, заповнюючи плейсхолдери даними договору.
    Додає таблицю товарів на місце плейсхолдера <товар>.

    Args:
        contract: Об'єкт Contract з даними для заповнення.
        template_path: Шлях до файлу шаблону (.docm або .docx).
        output_dir: Папка для збереження згенерованого файлу.

    Returns:
        Повний шлях до збереженого файлу або None у разі помилки.
    """
    if not os.path.exists(template_path):
        error_logger.log_error(FileNotFoundError(f"Шаблон не знайдено: {template_path}"), "Помилка генерації документа: Файл шаблону не знайдено")
        print(f"Помилка генерації: Файл шаблону не знайдено за шляхом: {template_path}")
        return None

    if not os.path.isdir(output_dir):
        error_logger.log_error(FileNotFoundError(f"Папка збереження не знайдена: {output_dir}"), "Помилка генерації документа: Папка збереження не знайдена")
        print(f"Помилка генерації: Папка для збереження не знайдена: {output_dir}")
        return None

    document = None
    try:
        document = Document(template_path)
        print(f"Відкрито шаблон: {template_path}")

        # Знаходимо параграф, де знаходиться плейсхолдер <товар>, якщо він є
        item_list_placeholder_paragraph = None
        item_list_placeholder_paragraph_parent = None
        item_list_placeholder_paragraph_index = None

        # Проходимо по всіх елементах body, щоб знайти параграф з плейсхолдером товарів
        for element in document.element.body.iter():
             # ПОПРАВКА ТУТ: Використовуємо element.tag == qn('w:p')
             if element.tag == qn('w:p'): # Це XML елемент параграфа
                  para = Paragraph(element, document._body)
                  if constants.ITEM_LIST_PLACEHOLDER in para.text:
                       item_list_placeholder_paragraph = para
                       item_list_placeholder_paragraph_parent = element.getparent()
                       item_list_placeholder_paragraph_index = item_list_placeholder_paragraph_parent.index(element)
                       print(f"Знайдено плейсхолдер товарів.")
                       break # Припускаємо, що плейсхолдер товарів зустрічається один раз в основному тексті


        # --- Заповнення стандартних плейсхолдерів ---
        # Ітеруємо по всіх елементах документа, які можуть містити текст (параграфи, комірки таблиць)
        # Використовуємо .iter() для наскрізного проходу по XML елементах
        for element in document.element.body.iter():
             # Шукаємо параграфи та комірки
             # ПОПРАВКА ТУТ: Використовуємо element.tag == qn('w:p')
             if element.tag == qn('w:p'): # Це параграф
                  para = Paragraph(element, document._body)
                  # Пропускаємо параграф з плейсхолдером товарів на цьому етапі
                  # (він вже обробляється окремо або буде видалений)
                  if para == item_list_placeholder_paragraph:
                       continue

                  for placeholder, value in contract.fields.items():
                      if placeholder not in [constants.ITEM_LIST_PLACEHOLDER]:
                           # Викликаємо нашу допрацьовану функцію заміни
                           _replace_placeholder_in_element(para, placeholder, str(value))

             # ПОПРАВКА ТУТ: Використовуємо element.tag == qn('w:tc')
             elif element.tag == qn('w:tc'): # Це комірка таблиці
                  # Створюємо об'єкт комірки. None замість батьківського елемента row
                  # достатньо для доступу до параграфів всередині комірки.
                  cell = _Cell(element, None)
                  for para in cell.paragraphs:
                       for placeholder, value in contract.fields.items():
                            if placeholder not in [constants.ITEM_LIST_PLACEHOLDER]:
                                # Викликаємо нашу допрацьовану функцію заміни
                                _replace_placeholder_in_element(para, placeholder, str(value))



        # --- Вставка таблиці товарів ---
        if item_list_placeholder_paragraph_index is not None and item_list_placeholder_paragraph_parent is not None:
             print("Обробка таблиці товарів...")

             if contract.items:
                 # Створюємо об'єкт таблиці з даними товарів
                 # _create_items_table тепер створює таблицю в окремому тимчасовому документі
                 items_table = _create_items_table(contract.items, contract) # Передаємо contract для суми

                 if items_table:
                      # Отримуємо XML елемент створеної таблиці з тимчасового документа
                      table_element = items_table._element

                      # Створюємо копію елемента таблиці, щоб вставити її в наш основний документ
                      # Використовуємо deepcopy для надійної копії
                      copied_table_element = deepcopy(table_element)

                      # Вставляємо скопійований XML елемент таблиці
                      # перед XML елементом параграфа з плейсхолдером в основному документі
                      item_list_placeholder_paragraph_parent.insert(item_list_placeholder_paragraph_index, copied_table_element)

                      # Видаляємо XML елемент параграфа з плейсхолдером товарів
                      item_list_placeholder_paragraph_parent.remove(item_list_placeholder_paragraph._element)
                      print("Таблицю товарів успішно створено та вставлено.")
                 else:
                      # Це не повинно статись, якщо contract.items не порожній і _create_items_table коректна
                      # Але якщо сталось, просто видаляємо плейсхолдер
                      item_list_placeholder_paragraph_parent.remove(item_list_placeholder_paragraph._element)
                      print("Попередження: Список товарів не порожній, але таблиця не була створена. Плейсхолдер видалено.")


             else:
                 # Якщо список товарів порожній, просто видаляємо плейсхолдер <товар>
                 if item_list_placeholder_paragraph: # Перевіряємо, чи був знайдений параграф
                      item_list_placeholder_paragraph_parent.remove(item_list_placeholder_paragraph._element)
                      print("Список товарів порожній, плейсхолдер товарів видалено.")

        elif constants.ITEM_LIST_PLACEHOLDER in contract.fields and contract.items:
             # Якщо плейсхолдер <товар> був у contract.fields (знайдено в шаблоні),
             # але параграф з ним не знайдено в body документа.
             # Це може статись, якщо плейсхолдер був у колонтитулі або іншому місці, яке ми не перебрали.
             print(f"Попередження: Плейсхолдер товарів '{constants.ITEM_LIST_PLACEHOLDER}' знайдено в даних договору, але не знайдено відповідного параграфа в основному тексті шаблону. Таблиця не буде вставлена.")
             # TODO: Обробка плейсхолдера товарів в інших частинах документа (колонтитули тощо)


        # --- Визначення імені файлу для збереження ---
        # Наприклад, "Договір [Контрагент] [Товар]".
        # Спробуємо взяти назву з поля <контрагент>, якщо воно є.
        # Або використати назву договору з об'єкта Contract.
        contract_name_for_filename = contract.fields.get("<контрагент>", contract.name)
        # Прибираємо некоректні символи для імені файлу
        contract_name_for_filename = re.sub(r'[<>:"/\\|?*]', '_', contract_name_for_filename)
        contract_name_for_filename = contract_name_for_filename.strip()
        if not contract_name_for_filename:
             contract_name_for_filename = "Без_назви_договору"


        # Приклад формату імені: "Договір [Назва_Договору]"
        filename = f"Договір {contract_name_for_filename}.docx"
        output_filepath = os.path.join(output_dir, filename)

        # Перевіряємо, чи файл з таким ім'ям вже існує, і додаємо суфікс, якщо потрібно
        base, ext = os.path.splitext(output_filepath)
        counter = 1
        while os.path.exists(output_filepath):
            output_filepath = f"{base}_{counter}{ext}"
            counter += 1


        # --- Збереження документа ---
        document.save(output_filepath)
        print(f"Документ успішно згенеровано: {output_filepath}")

        return output_filepath

    except Exception as e:
        error_logger.log_error(e, f"Критична помилка при генерації документа для договору '{contract.name}'")
        print(f"Критична помилка при генерації документа: {e}")
        return None


# TODO: Доопрацювати _replace_placeholder_in_element для ідеальної надійності збереження форматування при заміні через run-и.
# TODO: Можливо, додати підтримку багаторядкових полів за допомогою Text?
# Якщо використовується Text у GUI, то його значення може містити символи нового рядка.
# При заміні в Word, це потрібно конвертувати в розриви рядків (<w:br>) або нові параграфи.
# TODO: Додати підтримку плейсхолдерів у колонтитулах, зносках тощо (потрібно ітерувати не тільки body).
# TODO: Додати більше налаштувань форматування таблиці товарів (відступи, шрифти тощо).


# Приклад використання (для тестування)
if __name__ == "__main__":
    # Створюємо тимчасовий файл шаблону для тестування
    test_template_path = "temp_test_template.docm" # Можна також використовувати .docx

    try:
        # Створюємо простий документ Word з плейсхолдерами
        test_document = Document()
        test_document.add_paragraph("Це шаблонний документ.")
        test_document.add_paragraph("Назва заходу: <назва заходу>")
        test_document.add_paragraph("Адреса: <адреса>")
        test_document.add_paragraph("Дата: <дата>").add_run(" (Провести до <термін дії>)").bold = True # Плейсхолдер в run з форматуванням
        test_document.add_paragraph("Номер договору: <номер договору>")
        # Тест плейсхолдера через run-и:
        para_with_runs = test_document.add_paragraph("Контр")
        run1 = para_with_runs.add_run("аг")
        run2 = para_with_runs.add_run("ент:")
        run3 = para_with_runs.add_run(" <конт")
        run4 = para_with_runs.add_run("раген")
        run5 = para_with_runs.add_run("т>")
        run1.bold = True # Робимо частину жирною
        run5.italic = True # Робимо частину курсивом


        test_document.add_paragraph("Список товарів:")
        test_document.add_paragraph(constants.ITEM_LIST_PLACEHOLDER) # Плейсхолдер для таблиці товарів
        test_document.add_paragraph("Сума разом: " + constants.TOTAL_SUM_PLACEHOLDER) # Плейсхолдер суми разом
        test_document.add_paragraph("Сума прописом: " + constants.TEXT_SUM_PLACEHOLDER) # Плейсхолдер суми прописом
        test_document.add_paragraph("Інші умови: <інші умови>")

        # Додаємо плейсхолдер в таблицю для тестування заміни в таблицях
        table = test_document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Плейсхолдер в таблиці: <плейсхолдер в таблиці>"
        table.cell(0, 1).text = "Інше поле в таблиці: <інше поле в таблиці>"
        table.cell(1, 0).text = "Ще один плейсхолдер: <ще один плейсхолдер>"
        table.cell(1, 1).text = "Кінець таблиці."


        test_document.save(test_template_path)
        print(f"Створено тимчасовий шаблон: {test_template_path}")

        # Тестування пошуку плейсхолдерів
        found_placeholders = find_placeholders_in_template(test_template_path)
        print(f"\nРезультат пошуку плейсхолдерів: {found_placeholders}")

        # Створюємо тестовий об'єкт Contract
        test_contract = Contract("Тестовий договір для генерації")
        test_contract.fields = {
            "<назва заходу>": "Тестовий Фестиваль",
            "<адреса>": "м. Київ, вул. Тестова, 1",
            "<дата>": "2025-05-18",
            "<термін дії>": "31.12.2025", # Значення для плейсхолдера з форматуванням
            "<номер договору>": "ТД-123",
            "<контрагент>": "ТОВ \"Тестовий Контрагент\" з довгою назвою",
            "<імя представника>": "Іванов Іван Іванович",
            "<посада представника>": "Директор",
            "<підстава діяльності>": "на підставі Наказу № 5",
            "<реквізити>": "Р/Р 2600...",
            # Значення для <товар> не зберігається тут, він замінюється таблицею
            constants.TOTAL_SUM_PLACEHOLDER: "12345,67", # Ці значення зазвичай розраховуються в GUI
            constants.TEXT_SUM_PLACEHOLDER: "Дванадцять тисяч триста сорок п'ять гривень 67 копійок",
            "<інші умови>": "Тестові умови для договору.",
            "<плейсхолдер в таблиці>": "Замінено текст в таблиці!",
            "<інше поле в таблиці>": "Ще одне замінене поле",
            "<ще один плейсхолдер>": "І останнє поле в таблиці",
        }

        # Додаємо тестові товари до договору (з дробовими значеннями)
        test_contract.items.append(Item("Тестовий товар А (довга назва)", "11.11.11", 10.0, 100.0, 1000.0))
        test_contract.items.append(Item("Тестовий товар Б", "222-22", 5.5, 200.50, 1102.75))
        test_contract.items.append(Item("Тестовий товар В", "333", 1.0, 500.0, 500.0))
        test_contract.items.append(Item("Ще один товар", "444", 20.0, 10.0, 200.0))

        # Приклад договору без товарів
        test_contract_no_items = Contract("Договір без товарів")
        test_contract_no_items.fields = test_contract.fields.copy() # Копіюємо поля, але без товарів
        test_contract_no_items.items = [] # Пустий список товарів


        # Створюємо тимчасову папку для збереження
        test_output_dir = "temp_generated_docs"
        os.makedirs(test_output_dir, exist_ok=True)
        print(f"\nСтворено тимчасову папку для збереження: {test_output_dir}")


        # Тестування генерації документа З ТОВАРАМИ
        print("\nТестування генерації документа З ТОВАРАМИ...")
        generated_filepath_with_items = generate_document(test_contract, test_template_path, test_output_dir)

        if generated_filepath_with_items:
            print(f"Тестовий документ З ТОВАРАМИ згенеровано за шляхом: {generated_filepath_with_items}")
            print("БУДЬ ЛАСКА, ВІДКРИЙТЕ ЙОГО І ПЕРЕВІРТЕ!")
        else:
            print("Не вдалося згенерувати тестовий документ З ТОВАРАМИ.")

        # Тестування генерації документа БЕЗ ТОВАРІВ
        print("\nТестування генерації документа БЕЗ ТОВАРІВ...")
        generated_filepath_no_items = generate_document(test_contract_no_items, test_template_path, test_output_dir)

        if generated_filepath_no_items:
             print(f"Тестовий документ БЕЗ ТОВАРІВ згенеровано за шляхом: {generated_filepath_no_items}")
             print("БУДЬ ЛАСКА, ВІДКРИЙТЕ ЙОГО І ПЕРЕВІРТЕ (чи видалився плейсхолдер товарів).")
        else:
             print("Не вдалося згенерувати тестовий документ БЕЗ ТОВАРІВ.")


    except Exception as e:
        print(f"Виникла помилка під час тестового запуску document_processor: {e}")
        error_logger.log_error(e, "Помилка під час тестового запуску document_processor")

    finally:
        # Очистка тимчасових файлів (опціонально, залиште для перевірки)
        # if os.path.exists(test_template_path):
        #     os.remove(test_template_path)
        # if os.path.exists(test_output_dir) and os.path.isdir(test_output_dir):
        #      try:
        #           for f in os.listdir(test_output_dir):
        #                os.remove(os.path.join(test_output_dir, f))
        #           os.rmdir(test_output_dir)
        #      except Exception as e:
        #           print(f"Помилка при очистці тимчасової папки: {e}")

        pass # Залишаємо файли для перевірки після тесту